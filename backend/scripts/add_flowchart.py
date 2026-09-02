"""Add a visual flowchart to the existing Word document."""

import os
import sys
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════════════════════
# STEP 1: Generate flowchart image with matplotlib
# ═══════════════════════════════════════════════════════════════════

def create_flowchart(output_path):
    """Create a professional 4-layer pipeline flowchart."""
    fig, ax = plt.subplots(figsize=(11, 14))
    fig.patch.set_facecolor("#0F172A")
    ax.set_facecolor("#0F172A")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 16)
    ax.axis("off")

    # Title
    ax.text(5, 15.4, "PULSE AI DISCOVERY ENGINE", fontsize=20, fontweight="bold",
            ha="center", va="center", color="#818CF8", fontfamily="sans-serif")
    ax.text(5, 15.0, "End-to-End Pipeline Flowchart", fontsize=12,
            ha="center", va="center", color="#94A3B8", fontfamily="sans-serif")

    def draw_box(x, y, w, h, text, bg_color, border_color, text_color="white", fontsize=9, bold=False):
        box = FancyBboxPatch((x - w/2, y - h/2), w, h,
                             boxstyle="round,pad=0.12", linewidth=1.8,
                             edgecolor=border_color, facecolor=bg_color, zorder=2)
        ax.add_patch(box)
        fw = "bold" if bold else "normal"
        ax.text(x, y, text, fontsize=fontsize, ha="center", va="center",
                color=text_color, fontweight=fw, fontfamily="sans-serif",
                linespacing=1.35, zorder=3)

    def draw_arrow(x1, y1, x2, y2, color="#4B5563"):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                     arrowprops=dict(arrowstyle="-|>", color=color, lw=1.8,
                                     connectionstyle="arc3,rad=0"), zorder=1)

    def draw_label(x, y, text, color):
        ax.text(x, y, text, fontsize=11, fontweight="bold", ha="left", va="center",
                color=color, fontfamily="sans-serif",
                bbox=dict(boxstyle="round,pad=0.3", facecolor=color + "18", edgecolor=color, linewidth=1.2))

    # ── LAYER 1: INGESTION ──
    draw_label(0.3, 14.1, "LAYER 1 — Multi-Channel Ingestion", "#06B6D4")

    sources = [
        (1.2, 13.2, "Google Play\nStore"),
        (3.1, 13.2, "Apple App\nStore"),
        (5.0, 13.2, "Reddit"),
        (6.9, 13.2, "YouTube"),
        (8.8, 13.2, "Manual\nUpload"),
    ]
    for sx, sy, st in sources:
        draw_box(sx, sy, 1.5, 0.7, st, "#0E2A3A", "#06B6D4", fontsize=8)

    for sx, sy, _ in sources:
        draw_arrow(sx, sy - 0.35, 5.0, 12.25)

    draw_box(5.0, 12.0, 4.5, 0.55, "Normalize Text  →  SHA-256 Deduplicate  →  Enrich Metadata", "#0E2A3A", "#06B6D4", fontsize=8.5)
    draw_arrow(5.0, 11.7, 5.0, 11.15)
    draw_box(5.0, 10.9, 2.8, 0.45, "1,938 Documents in Database", "#0A2A1E", "#06B6D4", fontsize=8.5, bold=True)

    # ── LAYER 2: EXTRACTION ──
    draw_label(0.3, 10.1, "LAYER 2 — LLM Signal Extraction", "#8B5CF6")

    draw_arrow(5.0, 10.65, 5.0, 9.75)
    draw_box(5.0, 9.5, 3.8, 0.55, "Batch 20 Documents → Gemini Flash API\n(Structured JSON, temp=0.1)", "#1E153A", "#8B5CF6", fontsize=8.5)

    draw_arrow(5.0, 9.2, 5.0, 8.6)
    draw_box(5.0, 8.35, 3.5, 0.55, "Verbatim Quote Guardrail\n(Exact substring match vs. source text)", "#2D0A0A", "#EF4444", fontsize=8.5)

    draw_arrow(3.25, 8.1, 2.2, 7.65)
    draw_arrow(6.75, 8.1, 7.8, 7.65)

    draw_box(2.2, 7.4, 1.8, 0.45, "✅ Valid Signal\n→ Save to DB", "#0A2A1E", "#10B981", fontsize=7.5)
    draw_box(7.8, 7.4, 1.8, 0.45, "❌ Hallucinated\n→ Discard", "#2D0A0A", "#EF4444", fontsize=7.5)

    draw_arrow(2.2, 7.15, 5.0, 6.75)
    draw_box(5.0, 6.5, 2.8, 0.45, "1,554 Verified Extractions", "#1E153A", "#8B5CF6", fontsize=8.5, bold=True)

    # ── LAYER 3: AGGREGATION ──
    draw_label(0.3, 5.85, "LAYER 3 — Aggregation & Scoring", "#10B981")

    draw_arrow(5.0, 6.25, 5.0, 5.55)
    draw_box(5.0, 5.3, 3.6, 0.5, "Step 1: TF-IDF + SVD Embedding\n(100-dim dense vectors)", "#0A2A1E", "#10B981", fontsize=8.5)

    draw_arrow(5.0, 5.05, 5.0, 4.55)
    draw_box(5.0, 4.3, 3.6, 0.5, "Step 2: Agglomerative Clustering\n(cosine distance, silhouette-optimized k)", "#0A2A1E", "#10B981", fontsize=8.5)

    draw_arrow(5.0, 4.05, 5.0, 3.55)
    draw_box(5.0, 3.3, 3.6, 0.5, "Step 3: Taxonomy Labels + Consolidation\n(12-theme keyword matching → 8 unique)", "#0A2A1E", "#10B981", fontsize=8.5)

    draw_arrow(5.0, 3.05, 5.0, 2.55)
    draw_box(5.0, 2.3, 3.8, 0.5, "Step 4: 5-Dimension Composite Scoring\n(Freq + Tri + Conv + Seg + Act)", "#0A2A1E", "#10B981", fontsize=8.5)

    draw_arrow(5.0, 2.05, 5.0, 1.45)
    draw_box(5.0, 1.2, 2.8, 0.45, "8 Ranked Opportunity Areas", "#0A2A1E", "#10B981", fontsize=8.5, bold=True)

    # ── LAYER 4: SURFACE ──
    draw_label(0.3, 0.55, "LAYER 4 — Intelligence Surface", "#F59E0B")

    draw_arrow(3.5, 0.95, 2.5, 0.2)
    draw_arrow(6.5, 0.95, 7.5, 0.2)

    draw_box(2.5, -0.1, 2.5, 0.5, "Next.js Dashboard\n(Opportunity Table, Evidence)", "#2A200A", "#F59E0B", fontsize=8.5)
    draw_box(7.5, -0.1, 2.5, 0.5, "AI Insight Q&A\n(RAG + Gemini Synthesis)", "#2A200A", "#F59E0B", fontsize=8.5)

    plt.tight_layout()
    plt.savefig(output_path, dpi=200, bbox_inches="tight", facecolor="#0F172A", pad_inches=0.3)
    plt.close()
    print(f"Flowchart image saved: {output_path}")


# ═══════════════════════════════════════════════════════════════════
# STEP 2: Insert flowchart into existing Word document
# ═══════════════════════════════════════════════════════════════════

def insert_flowchart_into_docx(docx_path, image_path):
    """Insert the flowchart image after the Executive Overview in the Word doc."""
    doc = Document(docx_path)

    # Find the paragraph after "Key Capabilities" bullets and before the page break
    # We want to insert after Section 1 content (Executive Overview)
    insert_index = None
    for i, para in enumerate(doc.paragraphs):
        if "2. System Architecture" in para.text and para.style.name.startswith("Heading"):
            insert_index = i
            break

    if insert_index is None:
        # Fallback: insert after first page break / before section 2
        for i, para in enumerate(doc.paragraphs):
            if "System Architecture" in para.text:
                insert_index = i
                break

    if insert_index is None:
        insert_index = 10  # Safe fallback

    # We need to add content BEFORE the "2. System Architecture" heading
    # Insert a heading and the image before it
    body = doc.element.body

    # Create the heading element
    from docx.oxml.ns import qn
    from lxml import etree

    target_element = doc.paragraphs[insert_index]._element

    # Create "Pipeline Flowchart" heading paragraph
    heading_para = doc.add_heading("Pipeline Flowchart", level=2)
    for run in heading_para.runs:
        run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)

    # Move heading before target
    body.remove(heading_para._element)
    target_element.addprevious(heading_para._element)

    # Create image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(image_path, width=Inches(5.8))

    # Move image paragraph after heading
    body.remove(img_para._element)
    heading_para._element.addnext(img_para._element)

    # Add caption
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = caption_para.add_run("Figure 1: End-to-End Pipeline Flowchart — from 5 source channels through 4 processing layers")
    cr.font.size = Pt(8.5)
    cr.font.italic = True
    cr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    body.remove(caption_para._element)
    img_para._element.addnext(caption_para._element)

    # Add page break after caption
    pb_para = doc.add_paragraph()
    pb_run = pb_para.add_run()
    pb_run.add_break(docx.enum.text.WD_BREAK.PAGE)
    body.remove(pb_para._element)
    caption_para._element.addnext(pb_para._element)

    doc.save(docx_path)
    print(f"Flowchart inserted into: {docx_path}")


if __name__ == "__main__":
    import docx.enum.text

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    project_root = os.path.dirname(base_dir)

    image_path = os.path.join(base_dir, "scripts", "flowchart_pipeline.png")
    docx_path = os.path.join(project_root, "Pulse_AI_Discovery_Engine_Workflow.docx")

    # Step 1: Generate the flowchart image
    create_flowchart(image_path)

    # Step 2: Insert into Word document
    insert_flowchart_into_docx(docx_path, image_path)

    # Cleanup temp image
    # os.remove(image_path)

    print("\nDone! Flowchart added to Word document.")
