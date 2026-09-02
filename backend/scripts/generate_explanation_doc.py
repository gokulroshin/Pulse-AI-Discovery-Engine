import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Set cell padding in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    """Add a professional styled heading."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    
    # Custom color palette
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(30, 41, 59)     # Deep Slate
        run.font.size = Pt(18)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = RGBColor(79, 70, 229)    # Indigo Accent
        run.font.size = Pt(14)
        run.bold = True
    elif level == 3:
        run.font.color.rgb = RGBColor(15, 118, 110)   # Teal Accent
        run.font.size = Pt(12)
        run.bold = True
    return h

def add_callout_box(doc, title, body_text, fill_hex="F8FAFC", border_color="6366F1"):
    """Create a beautiful callout box with a colored left accent border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    
    # Left accent border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"💡 {title}\n" if title else "")
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = RGBColor(30, 41, 59)
    
    run_body = p.add_run(body_text)
    run_body.font.size = Pt(10)
    run_body.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def build_word_document(output_path: str):
    doc = docx.Document()
    
    # Page setup - Standard Letter / 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header and Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Pulse AI Discovery Engine — Plain English Guide")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(148, 163, 184)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Myntra Product & Growth | Confidential & Internal Guide")
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(148, 163, 184)
    
    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    
    tag_run = title_p.add_run("PRODUCT MANAGEMENT & GROWTH INTELLIGENCE\n")
    tag_run.font.size = Pt(9.5)
    tag_run.bold = True
    tag_run.font.color.rgb = RGBColor(79, 70, 229)
    
    title_run = title_p.add_run("How the Pulse AI Discovery Engine Works\n")
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    
    subtitle_run = title_p.add_run("A Simple, Non-Technical Guide to Explaining the System to Anyone")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Section 1: The 30-Second Elevator Pitch
    add_styled_heading(doc, "1. The 30-Second Pitch (Explain It in 3 Sentences)", level=1)
    
    add_callout_box(
        doc,
        "The Quick Summary",
        "Millions of fashion shoppers save clothes to their Myntra wishlist, but most never buy them within 30 days. "
        "Pulse is an automated AI engine that reads thousands of real customer discussions across Reddit, YouTube, and App Store reviews "
        "to pinpoint the exact reasons why shoppers hesitate (like confusing size charts or missing outfit styling ideas) "
        "and prioritizes exactly what Myntra should fix first—without giving discounts.",
        fill_hex="EEF2FF",
        border_color="4F46E5"
    )
    
    p = doc.add_paragraph()
    p.add_run("When someone asks you what this project is, you can say:\n")
    
    bp1 = doc.add_paragraph(style='List Bullet')
    bp1_r1 = bp1.add_run("The Problem: ")
    bp1_r1.bold = True
    bp1.add_run("Users show clear intent by wishlisting items, but get stuck before buying. Giving discounts hurts profit margins, so we must solve the real psychological and practical blockers.")
    
    bp2 = doc.add_paragraph(style='List Bullet')
    bp2_r1 = bp2.add_run("The AI Solution: ")
    bp2_r1.bold = True
    bp2.add_run("Instead of guessing why users drop off, our engine continuously listens to authentic customer conversations across 5 channels, extracts specific friction points, and ranks the highest-impact product opportunities.")
    
    bp3 = doc.add_paragraph(style='List Bullet')
    bp3_r1 = bp3.add_run("The Business Outcome: ")
    bp3_r1.bold = True
    bp3.add_run("Gives Product Managers clear, data-backed product roadmaps backed by 100% real, verifiable customer quotes.")
    
    # Section 2: Real-World Analogy
    add_styled_heading(doc, "2. The 'Super-Smart Research Intern' Analogy", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "To make this instantly click for any interviewer, executive, or teammate, use this simple analogy:"
    )
    
    add_callout_box(
        doc,
        "The Research Intern Analogy",
        "Imagine hiring a research intern and asking them to:\n"
        "1. Read through 2,000 customer reviews on the Google Play Store, Apple App Store, Reddit fashion forums, and YouTube haul comments.\n"
        "2. Use a highlighter pen to pull out every sentence where someone explains WHY they didn't buy or hesitated.\n"
        "3. Sort those 1,500+ highlighted sentences into sticky-note buckets on a whiteboard (e.g., 'Size Confusion', 'Return Worries', 'Can't Picture the Outfit').\n"
        "4. Calculate a priority score for each bucket so leadership knows which problem will unlock the most purchases.\n\n"
        "Pulse does all of this automatically in seconds, with zero human bias and complete mathematical scoring.",
        fill_hex="F0FDF4",
        border_color="16A34A"
    )
    
    # Section 3: The 4-Step Journey
    add_styled_heading(doc, "3. Step-by-Step: How It Works Under the Hood", level=1)
    
    p = doc.add_paragraph()
    p.add_run("The engine works in four clear, sequential stages:")
    
    # Step 1
    add_styled_heading(doc, "Step 1: Multi-Channel Ingestion (Listening Everywhere)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Customers don't just leave reviews on the app; they talk honestly in fashion communities on Reddit, complain on Twitter, and comment on YouTube try-on videos. "
        "Pulse gathers all these unstructured public conversations in one central database."
    )
    
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run("Channels Included: ").bold = True
    bp.add_run("Reddit (r/IndianFashionAddicts, r/Myntra), Google Play Store, Apple App Store, YouTube fashion haul comments, and E-commerce product reviews.")
    
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run("Smart Tagging: ").bold = True
    bp.add_run("Every review is automatically tagged with product category (Ethnic Wear, Western, Footwear), gender context (Women, Men), and brand price tier (Value, Mid, Premium).")
    
    # Step 2
    add_styled_heading(doc, "Step 2: Objective AI Extraction (Pulling Real Quotes)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Next, Google Gemini AI reads every single review and isolates the exact cause of hesitation along with the exact customer quote."
    )
    
    add_callout_box(
        doc,
        "Crucial Design Rule: Zero Confirmation Bias",
        "We do NOT tell the AI 'Find sizing problems' or 'Look for return issues'. "
        "The prompt is completely neutral: 'Extract the objective reason and the verbatim sentence without business bias.' "
        "This ensures our insights reflect what customers actually care about, rather than confirming what we already assumed.",
        fill_hex="FFFBEB",
        border_color="D97706"
    )
    
    # Step 3
    add_styled_heading(doc, "Step 3: Thematic Clustering (Grouping Similar Problems)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Instead of reading 1,500 separate quotes one by one, the engine uses mathematical clustering (Cosine Similarity) "
        "to group related problems into clear Opportunity Areas. For example, complaints about 'tight armholes in kurtis', 'inaccurate waist charts', and 'unpredictable brand sizing' "
        "are merged into one high-level theme: 'Fit & Sizing Confidence Gap'."
    )
    
    # Step 4
    add_styled_heading(doc, "Step 4: Business Opportunity Scoring (Prioritizing What to Build)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Not all customer complaints are equally important. Some are rare one-offs, while others block thousands of purchases. "
        "Pulse scores every opportunity using a 5-dimensional formula:"
    )
    
    # Table of Scoring Metrics
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Metric", "What It Measures in Plain English", "Weight"]
    col_widths = [Inches(1.8), Inches(3.7), Inches(1.0)]
    
    # Header row formatting
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)
    
    metrics_data = [
        ("1. Frequency", "How often do shoppers bring up this specific issue across the whole corpus?", "15%"),
        ("2. Triangulation", "Is this confirmed across multiple different platforms (e.g. Reddit AND Play Store)?", "25%"),
        ("3. Conversion Link", "How directly does this problem stop someone from clicking 'Buy Now' within 30 days?", "25%"),
        ("4. Segment Breadth", "Does this issue affect many product types (kurtis, jeans, dresses, shoes) or just one niche?", "15%"),
        ("5. Actionability", "Can Myntra product/design teams fix this with UX/features without relying on discounts?", "20%")
    ]
    
    for row_idx, (m_name, m_desc, m_wt) in enumerate(metrics_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate([m_name, m_desc, m_wt]):
            row_cells[c_idx].text = val
            row_cells[c_idx].width = col_widths[c_idx]
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
            p = row_cells[c_idx].paragraphs[0]
            p.runs[0].font.size = Pt(9.5)
            if c_idx == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
            elif c_idx == 2:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(79, 70, 229)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # Section 4: The 3 Core Discoveries
    add_styled_heading(doc, "4. The Top 3 Discoveries Uncovered by the Engine", level=1)
    
    p = doc.add_paragraph()
    p.add_run("When the engine analyzed the 1,938 customer reviews, three massive opportunity areas rose to the top:")
    
    # Discovery 1
    add_styled_heading(doc, "Discovery #1: Styling & Outfit Context Deficit (Rank #1 — Score: 0.87)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "• The Problem: Shoppers love an individual item (e.g. a kurti or skirt), but hesitate because they don't know what pants, footwear, or accessories to pair it with.\n"
        "• Verbatim Customer Quote: \"Wishlisted 8 tops waiting to compare fabrics and see if any styling reels show up on Instagram.\"\n"
        "• Product Opportunity: Add an AI 'Complete the Look' widget or community outfit pairings right inside the wishlist."
    )
    
    # Discovery 2
    add_styled_heading(doc, "Discovery #2: Fit & Sizing Confidence Gap (Rank #3 — Score: 0.55)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "• The Problem: Standard size charts are generic and don't account for stretch, height, or brand variations. Shoppers go to YouTube haul videos just to see how it looks on someone of their height.\n"
        "• Verbatim Customer Quote: \"Bought 3 kurtas for Diwali, 2 had sizing off by at least 2 inches at chest. Sizing charts need true customer measurement photos.\"\n"
        "• Product Opportunity: Show verified customer photos filtered by the shopper's exact height and body type."
    )
    
    # Discovery 3
    add_styled_heading(doc, "Discovery #3: Studio Lighting vs. Real-Life Texture Mismatch", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "• The Problem: Heavy studio lighting in catalog photos hides actual fabric transparency and natural daylight color shades.\n"
        "• Verbatim Customer Quote: \"Looked at YouTube haul to check actual dupatta drape and embroidery shine because catalog photos are too edited.\"\n"
        "• Product Opportunity: Introduce unedited daylight customer photos and fabric drape video previews."
    )
    
    # Section 5: AI Insight Search
    add_styled_heading(doc, "5. The AI Insight Search (Conversational PM Research)", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "The engine also includes an interactive AI Search Bar that acts like a senior qualitative research analyst. "
        "A Product Manager can type any question in plain English, and the engine retrieves relevant customer reviews, links them to ranked opportunities, and generates an evidence-grounded answer."
    )
    
    add_callout_box(
        doc,
        "Example PM Query in the Search Bar",
        "PM Question: 'Why do shoppers hesitate before buying Ethnic Wear?'\n\n"
        "Engine Answer: 'In Ethnic Wear, shoppers save items weeks in advance for festive occasions, but pause checkout due to two main doubts: (1) tightness in kurti armholes/blouse stitches, and (2) whether the dupatta fabric is see-through or scratchy in natural daylight. Verified customer try-on photos significantly accelerate checkout.'",
        fill_hex="FDF4FF",
        border_color="C026D3"
    )
    
    # Section 6: Interview & Presentation FAQ
    add_styled_heading(doc, "6. Interview & Pitch Cheat Sheet (Common Questions & Winning Answers)", level=1)
    
    faq_items = [
        (
            "Q: Why did you build an automated engine instead of just asking ChatGPT?",
            "A: ChatGPT in a regular chat interface hallucinates, forgets context, and only gives generic advice. "
            "Pulse is a complete batch data pipeline that processes thousands of real reviews, extracts verifiable verbatim quotes, "
            "applies multi-dimensional mathematical scoring, and provides 100% evidence traceability back to the source URL and platform."
        ),
        (
            "Q: Why are discounts explicitly prohibited in this project?",
            "A: Giving a 10% coupon is an expensive, short-term band-aid that erodes Myntra's profit margins. "
            "Furthermore, discounts don't fix the root problem: if a shopper is afraid a kurti will be too tight or look see-through, "
            "a 10% discount won't make them confident. Fixing fit clarity and styling context permanently drives organic conversion."
        ),
        (
            "Q: What is 'Triangulation' and why is it so important?",
            "A: Triangulation means checking if a customer complaint appears across multiple independent sources. "
            "If someone complains about sizing on the Play Store, it could be an isolated gripe. But if users on Reddit, App Store, AND YouTube "
            "all complain about the exact same sizing ambiguity, that proves it's a systemic, high-priority product opportunity."
        ),
        (
            "Q: What is the tech stack behind Pulse?",
            "A: Backend is built with Python, FastAPI, and SQLAlchemy for high-speed data processing; "
            "Google Gemini 3.6 Flash for unbiased quote extraction and RAG synthesis; "
            "Scikit-Learn for agglomerative semantic clustering; "
            "and Next.js 15 + TypeScript + Vanilla CSS for the executive dashboard."
        )
    ]
    
    for q, a in faq_items:
        p = doc.add_paragraph()
        q_run = p.add_run(f"{q}\n")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(30, 41, 59)
        q_run.font.size = Pt(10.5)
        
        a_run = p.add_run(a)
        a_run.font.color.rgb = RGBColor(71, 85, 105)
        a_run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)
        
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    output_file = "d:/3. Career/Product Management/IDE/Myntra- Grad Project/AI Discovery Engine/Pulse_AI_Discovery_Engine_Explained_Simply.docx"
    build_word_document(output_file)
